def run_telegram_agent():
    import asyncio
    import threading
    import sys
    stop_event = threading.Event()


    def input_listener():
        while not stop_event.is_set():
            try:
                comando = input("[Escribe 'sal' para volver al menú principal]: ").strip().lower()
                if comando == 'sal':
                    while True:
                        confirm = input("¿Seguro que quieres salir del agente de Telegram? (si/no): ").strip().lower()
                        if confirm in ('si', 's'):
                            stop_event.set()
                            print("\n[Saliendo del agente de Telegram...]")
                            # Forzar cierre de asyncio
                            try:
                                asyncio.get_event_loop().stop()
                            except Exception:
                                pass
                            return
                        elif confirm in ('no', 'n'):
                            print("Continuando con el agente de Telegram.")
                            break
            except EOFError:
                break

    listener_thread = threading.Thread(target=input_listener, daemon=True)
    listener_thread.start()

    async def main_with_stop():
        task = asyncio.create_task(main())
        while not stop_event.is_set():
            await asyncio.sleep(0.5)
        # Cuando stop_event esté activo, desconectar el cliente
        try:
            from telethon import TelegramClient
            client = TelegramClient(SESSION_NAME, API_ID, API_HASH)
            await client.disconnect()
        except Exception:
            pass
        task.cancel()

    try:
        asyncio.run(main_with_stop())
    except Exception:
        pass
"""
Agente MCP para monitorizar Telegram usando el servidor MCP definido en .vscode/mcp.json
"""
import os
import json
import sys
import re
from rich.console import Console
from rich.panel import Panel
from telethon import TelegramClient, events
from url_processor import process_url
from docx import Document
from datetime import datetime
import asyncio
import signal
import pickle
import tempfile
import requests
from docx.shared import RGBColor

# Leer configuración MCP
CONFIG_PATH = os.path.join(os.path.dirname(__file__), '.vscode', 'mcp.json')
try:
    with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
        config = json.load(f)
        tg_env = config['servers']['telegram']['env']
        API_ID = tg_env['TG_APP_ID']
        API_HASH = tg_env['TG_API_HASH']
except Exception as e:
    print(f'Error leyendo configuración MCP: {e}')
    sys.exit(1)

GROUP_NAME = 'GROUP_ID'  # Cambia esto por el nombre o ID del grupo
SESSION_NAME = 'telegram_agent_session'

# Expresión regular para extraer URLs
def extract_urls(text):
    return re.findall(r'(https?://[^\s]+)', text)

console = Console()
async def main():
    client = TelegramClient(SESSION_NAME, API_ID, API_HASH)
    await client.start()
    console.print(Panel('[bold blue]Agente MCP iniciado. Escuchando mensajes...[/bold blue]', expand=False))

    mensajes_chat = []
    ids_procesados = set()
    ids_file = os.path.join('resumenes', 'ids_procesados.pkl')
    # Cargar IDs procesados si existen
    if os.path.exists(ids_file):
        with open(ids_file, 'rb') as f:
            ids_procesados = pickle.load(f)

    async def guardar_resumen_global():
        if not mensajes_chat:
            console.print('[yellow]No hay mensajes para resumir.[/yellow]')
            return
        texto_completo = '\n'.join(mensajes_chat)
        resumen_global = process_url('https://dummy.url')  # Usamos la función de resumen, pero solo para el texto
        resumen_global['title'] = 'Resumen global del chat'
        resumen_global['summary'] = texto_completo[:2000] + ('...' if len(texto_completo) > 2000 else '')
        doc = Document()
        doc.add_heading(resumen_global.get('title', 'Resumen'), 0)
        doc.add_paragraph(f"Resumen: {resumen_global.get('summary', '')}")
        filename = f"resumen_global_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join('resumenes', filename)
        doc.save(filepath)
        console.print(f'✅ [bold green]Resumen global guardado en {filepath}[/bold green]')

    # Documento único para todos los resúmenes
    # Obtener nombre del grupo
    try:
        grupo_entidad = await client.get_entity(GROUP_NAME)
        nombre_grupo = grupo_entidad.title if hasattr(grupo_entidad, 'title') else str(GROUP_NAME)
    except Exception:
        nombre_grupo = str(GROUP_NAME)

    doc_path = os.path.join('resumenes', f'resumenes_{nombre_grupo}.docx')
    if os.path.exists(doc_path):
        doc = Document(doc_path)
    else:
        doc = Document()
        doc.add_heading(f'Resúmenes del grupo: {nombre_grupo}', 0)

    @client.on(events.NewMessage(chats=GROUP_NAME))
    async def handler(event):
        message = event.message.message
        msg_id = event.message.id
        if msg_id in ids_procesados:
            return  # Ya procesado
        if not message:
            return
        mensajes_chat.append(message)
        ids_procesados.add(msg_id)
        with open(ids_file, 'wb') as f:
            pickle.dump(ids_procesados, f)
        urls = extract_urls(message)
        if urls:
            console.print(f'🔗 [cyan]URLs encontradas:[/cyan] {urls}')
            for url in urls:
                resumen = process_url(url)
                console.print(f'📝 [green]Resumen generado:[/green] {resumen}')
                # Emoticonos temáticos
                url_emoji = '🔗'
                resumen_emoji = '📝'
                desc_emoji = '💡'
                img_emoji = '🖼️'
                id_emoji = '🆔'

                # Título en azul oscuro
                title = resumen.get('title', 'Resumen')
                heading = doc.add_heading(f"{resumen_emoji} {title}", level=1)
                run = heading.runs[0]
                run.font.color.rgb = RGBColor(0, 32, 96)  # Azul oscuro RGB

                # Grupo (subtítulo amarillo oscuro)
                p_group = doc.add_paragraph()
                run_group = p_group.add_run(f"👥 Grupo: {nombre_grupo}")
                run_group.bold = True
                run_group.font.color.rgb = RGBColor(191, 144, 0)

                # ID mensaje (subtítulo amarillo oscuro)
                p_id = doc.add_paragraph()
                run_id = p_id.add_run(f"{id_emoji} ID mensaje: {msg_id}")
                run_id.bold = True
                run_id.font.color.rgb = RGBColor(191, 144, 0)  # Amarillo oscuro RGB

                # URL (subtítulo amarillo oscuro)
                p_url = doc.add_paragraph()
                run_url = p_url.add_run(f"{url_emoji} URL: {resumen.get('url', '')}")
                run_url.bold = True
                run_url.font.color.rgb = RGBColor(191, 144, 0)

                # Resumen (subtítulo amarillo oscuro o error)
                resumen_texto = resumen.get('summary', '')
                # Generar descripción breve (primeras 20 palabras del resumen, sin repetir el resumen completo)
                if resumen_texto.startswith('Error al procesar la URL:'):
                    # Mensaje de error destacado
                    p_res = doc.add_paragraph()
                    run_res = p_res.add_run(f"⚠️ Error al procesar la URL: {resumen.get('url', '')}")
                    run_res.bold = True
                    run_res.font.color.rgb = RGBColor(255, 0, 0)  # Rojo
                    p_desc = doc.add_paragraph()
                    run_desc = p_desc.add_run(f"❌ Motivo: {resumen_texto.replace('Error al procesar la URL:', '').strip()}")
                    run_desc.italic = True
                    run_desc.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    # Descripción breve ANTES del resumen
                    palabras = resumen_texto.split()
                    breve = ' '.join(palabras[:20]) + ('...' if len(palabras) > 20 else '')
                    if breve:
                        p_desc = doc.add_paragraph()
                        run_desc = p_desc.add_run(f"{desc_emoji} Descripción breve de la URL: {breve}")
                        run_desc.italic = True
                    p_res = doc.add_paragraph()
                    run_res = p_res.add_run(f"{resumen_emoji} Resumen: {resumen_texto}")
                    run_res.font.color.rgb = RGBColor(191, 144, 0)

                # Imagen (unificado)
                img_url = resumen.get('image')
                if img_url:
                    try:
                        response = requests.get(img_url, timeout=10)
                        response.raise_for_status()
                        from PIL import Image
                        from docx.shared import Inches
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp_img:
                            tmp_img.write(response.content)
                            tmp_img.flush()
                            # Calcular tamaño óptimo
                            with Image.open(tmp_img.name) as im:
                                width, height = im.size
                                max_width_in = 4.5  # Máximo ancho en pulgadas
                                max_height_in = 3.0  # Máximo alto en pulgadas
                                dpi = im.info.get('dpi', (96, 96))[0]
                                width_in = width / dpi
                                height_in = height / dpi
                                scale = min(max_width_in / width_in, max_height_in / height_in, 1.0)
                                final_width = Inches(width_in * scale)
                                final_height = Inches(height_in * scale)
                            # Añadir párrafo centrado para la imagen
                            p_img = doc.add_paragraph(f"{img_emoji} Imagen:")
                            p_img.alignment = 1
                            picture = doc.add_picture(tmp_img.name, width=final_width, height=final_height)
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = 1
                        os.unlink(tmp_img.name)
                    except Exception as e:
                        p_img = doc.add_paragraph()
                        run_img = p_img.add_run(f"{img_emoji} Imagen (no se pudo insertar): {img_url}")
                        run_img.font.color.rgb = RGBColor(255, 0, 0)
            doc.save(doc_path)
            console.print(f'✅ [bold green]Resumen guardado en {doc_path}[/bold green]')

    # Procesar historial completo al iniciar
    console.print('🔎 [yellow]Descargando historial completo del grupo...[/yellow]')
    async for message in client.iter_messages(GROUP_NAME, reverse=True):
        msg_id = message.id
        if msg_id in ids_procesados:
            continue
        if not message.message:
            continue  # Ignora mensajes sin texto
        mensajes_chat.append(message.message)
        ids_procesados.add(msg_id)
        with open(ids_file, 'wb') as f:
            pickle.dump(ids_procesados, f)
        urls = extract_urls(message.message)
        if urls:
            console.print(f'🔗 [cyan]URLs encontradas en historial:[/cyan] {urls}')
            for url in urls:
                resumen = process_url(url)
                console.print(f'📝 [green]Resumen generado:[/green] {resumen}')
                # Emoticonos temáticos
                url_emoji = '🔗'
                resumen_emoji = '📝'
                desc_emoji = '💡'
                img_emoji = '🖼️'
                id_emoji = '🆔'

                # Título en azul oscuro
                title = resumen.get('title', 'Resumen')
                heading = doc.add_heading(f"{resumen_emoji} {title}", level=1)
                run = heading.runs[0]
                run.font.color.rgb = RGBColor(0, 32, 96)  # Azul oscuro RGB

                # Grupo (subtítulo amarillo oscuro)
                p_group = doc.add_paragraph()
                run_group = p_group.add_run(f"👥 Grupo: {nombre_grupo}")
                run_group.bold = True
                run_group.font.color.rgb = RGBColor(191, 144, 0)

                # ID mensaje (subtítulo amarillo oscuro)
                p_id = doc.add_paragraph()
                run_id = p_id.add_run(f"{id_emoji} ID mensaje: {msg_id}")
                run_id.bold = True
                run_id.font.color.rgb = RGBColor(191, 144, 0)  # Amarillo oscuro RGB

                # URL (subtítulo amarillo oscuro)
                p_url = doc.add_paragraph()
                run_url = p_url.add_run(f"{url_emoji} URL: {resumen.get('url', '')}")
                run_url.bold = True
                run_url.font.color.rgb = RGBColor(191, 144, 0)

                # Resumen (subtítulo amarillo oscuro o error)
                resumen_texto = resumen.get('summary', '')
                if resumen_texto.startswith('Error al procesar la URL:'):
                    # Mensaje de error destacado
                    p_res = doc.add_paragraph()
                    run_res = p_res.add_run(f"⚠️ Error al procesar la URL: {resumen.get('url', '')}")
                    run_res.bold = True
                    run_res.font.color.rgb = RGBColor(255, 0, 0)  # Rojo
                    p_desc = doc.add_paragraph()
                    run_desc = p_desc.add_run(f"❌ Motivo: {resumen_texto.replace('Error al procesar la URL:', '').strip()}")
                    run_desc.italic = True
                    run_desc.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    p_res = doc.add_paragraph()
                    run_res = p_res.add_run(f"{resumen_emoji} Resumen: {resumen_texto}")
                    run_res.font.color.rgb = RGBColor(191, 144, 0)
                    # Descripción breve de la URL
                    if resumen_texto:
                        p_desc = doc.add_paragraph()
                        run_desc = p_desc.add_run(f"{desc_emoji} Descripción breve de la URL: {resumen_texto}")
                        run_desc.italic = True

                # Imagen (unificado)
                img_url = resumen.get('image')
                if img_url:
                    try:
                        response = requests.get(img_url, timeout=10)
                        response.raise_for_status()
                        from PIL import Image
                        from docx.shared import Inches
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp_img:
                            tmp_img.write(response.content)
                            tmp_img.flush()
                            # Calcular tamaño óptimo
                            with Image.open(tmp_img.name) as im:
                                width, height = im.size
                                max_width_in = 4.5  # Máximo ancho en pulgadas
                                max_height_in = 3.0  # Máximo alto en pulgadas
                                dpi = im.info.get('dpi', (96, 96))[0]
                                width_in = width / dpi
                                height_in = height / dpi
                                scale = min(max_width_in / width_in, max_height_in / height_in, 1.0)
                                final_width = Inches(width_in * scale)
                                final_height = Inches(height_in * scale)
                            # Añadir párrafo centrado para la imagen
                            p_img = doc.add_paragraph(f"{img_emoji} Imagen:")
                            p_img.alignment = 1
                            picture = doc.add_picture(tmp_img.name, width=final_width, height=final_height)
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = 1
                        os.unlink(tmp_img.name)
                    except Exception as e:
                        p_img = doc.add_paragraph()
                        run_img = p_img.add_run(f"{img_emoji} Imagen (no se pudo insertar): {img_url}")
                        run_img.font.color.rgb = RGBColor(255, 0, 0)
    doc.save(doc_path)
    console.print(f'✅ [bold green]Todos los resúmenes guardados en {doc_path}[/bold green]')

    # Tarea para resumen diario
    async def resumen_diario():
        while True:
            await asyncio.sleep(86400)  # 24 horas
            await guardar_resumen_global()

    # Manejar Ctrl+C para guardar resumen global
    def shutdown_handler():
        print("\nGenerando resumen global antes de salir...")
        loop = asyncio.get_event_loop()
        loop.run_until_complete(guardar_resumen_global())
        sys.exit(0)

    signal.signal(signal.SIGINT, lambda s, f: shutdown_handler())
    # Iniciar tarea de resumen diario
    asyncio.create_task(resumen_diario())
    await client.run_until_disconnected()

if __name__ == '__main__':
    import asyncio
    asyncio.run(main())
